"""
将pdf转换成word
"""
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document

document = Document()


def parse():
    # rb以二进制读模式打开本地pdf文件
    rpath = r'C:\Users\hdp\Desktop\transformationPaper_EN_CN'
    # doc_name = r'2-1-5021-distributed-representations-of-words-and-phrases-and-their-compositionality.pdf'
    #doc_name = r'[3]Reducing the dimensionality of data with neural networks_science.pdf'
    # doc_name = r'[2]_Hinton,_Geoffrey_E,_Simon_Osindero,_and_Yee-Whye_Teh_A_fast_learning_algorithm_for_deep_belief_nets_Neural_computation_187_(2006)__1527-1554_.pdf'
    #seq2seq
    doc_name = r'[36]_Sutskever,_Ilya,_Oriol_Vinyals,_and_Quoc_V_Le_Sequence_to_sequence_learning_with_neural_networks_Advances_in_neural_information_processing_systems_2014_.pdf'
    p = rpath+'\\'+doc_name
    print(p)
    fn = open(p,'rb')
    # 创建一个pdf文档分析器
    parser = PDFParser(fn)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)

    # 提供初始化密码doc.initialize("lianxipython")
    # 如果没有密码 就创建一个空的字符串
    doc.initialize("")
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed

    else:
        # 创建PDf资源管理器
        resource = PDFResourceManager()
        # 创建一个PDF参数分析器
        laparams = LAParams()
        # 创建聚合器,用于读取文档的对象
        device = PDFPageAggregator(resource, laparams=laparams)
        # 创建解释器，对文档编码，解释成Python能够识别的格式
        interpreter = PDFPageInterpreter(resource, device)
        # 循环遍历列表，每次处理一页的内容
        # doc.get_pages() 获取page列表
        doc
        for page in doc.get_pages():
            # 利用解释器的process_page()方法解析读取单独页数
            interpreter.process_page(page)
            # 使用聚合器get_result()方法获取内容
            layout = device.get_result()
            # 这里layout是一个LTPage对象,里面存放着这个page解析出的各种对象
            for out in layout:
                # 判断是否含有get_text()方法，获取我们想要的文字
                if hasattr(out, "get_text"):
                    # print(out.get_text(), type(out.get_text()))
                    content = out.get_text().replace(u'\xa0', u' ')  # 将'\xa0'替换成u' '空格，这个\xa0就是&nbps空格
                    # with open('test.txt','a') as f:
                    #     f.write(out.get_text().replace(u'\xa0', u' ')+'\n')
                    document.add_paragraph(
                        content, style= 'List Number'#'ListBullet'  # 添加段落，样式为unordered list类型
                    )
                document.save(rpath+'\\'+doc_name.replace('.pdf','.docx'))  # 保存这个文档


if __name__ == '__main__':
    parse()

